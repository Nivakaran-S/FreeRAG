"""Document loader for various file formats."""

from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Dict, Any


@dataclass
class Document:
    """Represents a loaded document."""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    @property
    def source(self) -> str:
        """Get document source path."""
        return self.metadata.get("source", "unknown")


class DocumentLoader:
    """Load documents from various file formats."""
    
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
    
    def __init__(self):
        """Initialize the document loader."""
        self._pdf_loader = None
        self._docx_loader = None
    
    def load_file(self, file_path: str) -> Document:
        """Load a single file.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            Loaded document.
            
        Raises:
            ValueError: If file format is not supported.
            FileNotFoundError: If file doesn't exist.
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )
        
        content = self._load_by_extension(path, extension)
        
        return Document(
            content=content,
            metadata={
                "source": str(path.absolute()),
                "filename": path.name,
                "extension": extension
            }
        )
    
    def load_directory(
        self, 
        directory_path: str,
        recursive: bool = True
    ) -> List[Document]:
        """Load all supported files from a directory.
        
        Args:
            directory_path: Path to the directory.
            recursive: Whether to search recursively.
            
        Returns:
            List of loaded documents.
        """
        path = Path(directory_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        if not path.is_dir():
            raise ValueError(f"Not a directory: {directory_path}")
        
        documents = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load_file(str(file_path))
                    documents.append(doc)
                    print(f"Loaded: {file_path.name}")
                except Exception as e:
                    print(f"Warning: Failed to load {file_path.name}: {e}")
        
        return documents
    
    def _load_by_extension(self, path: Path, extension: str) -> str:
        """Load file content based on extension.
        
        Args:
            path: File path.
            extension: File extension.
            
        Returns:
            File content as string.
        """
        if extension in {".txt", ".md"}:
            return self._load_text(path)
        elif extension == ".pdf":
            return self._load_pdf(path)
        elif extension == ".docx":
            return self._load_docx(path)
        else:
            raise ValueError(f"Unknown extension: {extension}")
    
    def _load_text(self, path: Path) -> str:
        """Load plain text file."""
        return path.read_text(encoding="utf-8")
    
    def _load_pdf(self, path: Path) -> str:
        """Load PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required for PDF files: pip install pypdf")
        
        reader = PdfReader(str(path))
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    def _load_docx(self, path: Path) -> str:
        """Load DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX files: pip install python-docx")
        
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
